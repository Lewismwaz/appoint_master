from bs4 import BeautifulSoup
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Spacer, Image
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

# Read the HTML content
with open('appoint_app/appointment_report.html', 'r') as file:
    html_content = file.read()

# Parse HTML
soup = BeautifulSoup(html_content, 'html.parser')

# Find table
table = soup.find('table')

# Extract table data into a list of lists
data = []
for row in table.find_all('tr'):
    row_data = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
    data.append(row_data)

# Logo path
logo_path = 'static/appoint_app/images/appoint-master-logo.png'

# Generate CSV
df = pd.DataFrame(data[1:], columns=data[0])
df.to_csv('report.csv', index=False)

# Generate Excel
df.to_excel('report.xlsx', index=False)

# Generate PDF
pdf = SimpleDocTemplate("report.pdf", pagesize=letter)
table_data = [data[0]] + [data[i] for i in range(1, len(data))]
pdf_table = Table(table_data)

# Add style to table
style = TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.gray),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)])
pdf_table.setStyle(style)

# Add logo and footer
elements = [Image(logo_path, width=100, height=50), Spacer(1, 20), pdf_table]

# Build PDF
pdf.build(elements)

# Generate DOCX
doc = Document()
# Add logo
doc.add_picture(logo_path, width=Inches(2))
doc.add_paragraph()

# Add table
table = doc.add_table(rows=len(data), cols=len(data[0]))
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.cell(i, j).text = cell

# Add footer
footer = doc.sections[0].footer
footer.paragraphs[0].text = "© 2024 Appoint Master, KUHC"

# Save DOCX
doc.save('report.docx')
